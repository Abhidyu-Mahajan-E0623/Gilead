from docx import Document

doc_path = r"c:\Users\AbhidyuMahajan\OneDrive - ProcDNA Analytics Pvt. Ltd\Desktop\Projects\Gilead_Demo\Input\GILEAD Field Inquiry Script v1.0.docx"
doc = Document(doc_path)

print(f"Num tables: {len(doc.tables)}")
for i, table in enumerate(doc.tables):
    print(f"\n--- TABLE {i} ---")
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"R{r}C{c}: {text[:80]}")
                if text.startswith("Resolution and Response to Rep:"):
                    print(f"  FULL TEXT OF CELL: {text}")
