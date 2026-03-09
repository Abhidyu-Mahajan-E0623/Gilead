from docx import Document
import json
import re

doc_path = r"c:\Users\AbhidyuMahajan\OneDrive - ProcDNA Analytics Pvt. Ltd\Desktop\Projects\Gilead_Demo\Input\GILEAD Field Inquiry Script v1.0.docx"
json_path = r"c:\Users\AbhidyuMahajan\OneDrive - ProcDNA Analytics Pvt. Ltd\Desktop\Projects\Gilead_Demo\Input\GILEAD_Field_Inquiry_Playbook.json"

doc = Document(doc_path)
resolutions = []

for table in doc.tables:
    if len(table.rows) >= 7:
        # Check if R6C0 is "Resolution & Response to Rep"
        header = table.rows[6].cells[0].text.strip()
        if "Resolution" in header:
            # Reconstruct the paragraph spacing by joining paragraph texts inside the cell
            cell = table.rows[6].cells[1]
            paras = []
            for p in cell.paragraphs:
                text = p.text.strip()
                if text:
                    # Remove the bullet character (often something like a symbol or a literal dot)
                    # The text might start with a special character like  or ·
                    text = re.sub(r'^[\W_]+\s*', '', text)
                    
                    # Convert em dashes as user requested
                    text = text.replace(" — ", " - ").replace("—", "-")
                    
                    paras.append(text)
            
            if paras:
                resolutions.append("\n".join(paras))

with open(json_path, 'r', encoding='utf-8') as f:
    data = json.load(f)

print(f"Found {len(resolutions)} resolutions in the Word tables.")

if len(resolutions) == len(data.get("inquiries", [])):
    for i, item in enumerate(data["inquiries"]):
        item["resolution_and_response_to_rep"] = resolutions[i]
        
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print("Successfully updated the JSON with Word document table formatting.")
else:
    print(f"Error: Found {len(resolutions)} resolutions, but JSON has {len(data.get('inquiries', []))} inquiries.")
    for i in range(min(5, len(resolutions))):
        print(f"--- Resolution {i+1} ---")
        print(resolutions[i])
